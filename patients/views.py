import io
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, logout
from django.contrib.auth.decorators import login_required
from .models import Patient
from .forms import RegisterForm, PatientSearchForm
from docx import Document
from django.http import HttpResponse
from django.contrib.auth.models import User
from django.contrib.auth import authenticate


def home(request):
    return render(request, 'patients/home.html')


def register_view(request):
    if request.method == 'POST':
        form = RegisterForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.first_name = form.cleaned_data['first_name']
            user.last_name = form.cleaned_data['last_name']
            user.save()
            login(request, user)
            return redirect('home')
    else:
        form = RegisterForm()
    return render(request, 'patients/register.html', {'form': form})


def login_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('home')
    return render(request, 'patients/login.html')


@login_required
def my_patients(request):
    patients = Patient.objects.filter(attending_doctor=request.user)
    return render(request, 'patients/my_patients.html', {'patients': patients})


@login_required
def print_doc(request):
    patients = []
    form = PatientSearchForm(request.GET or None)

    if form.is_valid() and form.cleaned_data['full_name']:
        patients = Patient.objects.filter(full_name__icontains=form.cleaned_data['full_name'])

    # Показываем только тех пациентов, где лечащий врач == текущий пользователь
    patients_list = Patient.objects.filter(attending_doctor=request.user)

    return render(request, 'patients/print_doc.html', {
        'form': form,
        'patients': patients,
        'patients_list': patients_list
    })


def generate_recommendation(request, patient_id):
    patient = get_object_or_404(Patient, id=patient_id)

    doc = Document()
    doc.add_heading("Рекомендации по реабилитации", level=1)

    doc.add_paragraph(f"ФИО: {patient.full_name}")

    if patient.birth_date:
        age = patient.get_age()
        birth_date_str = patient.birth_date.strftime("%d.%m.%Y")
        doc.add_paragraph(f"Дата рождения: {birth_date_str} ({age} лет)")
    else:
        doc.add_paragraph("Дата рождения: не указана")

    doc.add_paragraph(f"Пол: {patient.get_gender_display() or 'Не указан'}")
    doc.add_paragraph(f"Диагноз: {patient.diagnosis}")
    doc.add_paragraph(f"Рекомендации: {patient.recommendations}")
    doc.add_paragraph(f"Лечащий врач: {request.user.get_full_name() or request.user.username}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=recommendation_{patient.id}.docx'
    return response


def logout_view(request):
    logout(request)
    return redirect('home')